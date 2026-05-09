from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document

from .models import ExtractionModels
from .text_utils import clean_text


GLINER_LABELS_DE = [
    "Zielgruppe",
    "Projektmaßnahme",
    "Projektziel",
    "Wirkung",
    "Region",
    "Kooperationspartner",
]

RELEVANT_TERMS = [
    "zielgruppe",
    "teilnehm",
    "maßnahme",
    "projektinhalt",
    "ziel",
    "wirkung",
    "beratung",
    "kurs",
    "praktika",
    "lehr",
    "region",
    "standort",
]


def docx_text_blocks(path: Path) -> list[dict]:
    doc = Document(path)
    blocks = []

    for idx, para in enumerate(doc.paragraphs):
        text = clean_text(para.text)
        if text:
            blocks.append(
                {
                    "source_file": str(path),
                    "kind": "paragraph",
                    "index": idx,
                    "style": para.style.name,
                    "text": text,
                }
            )

    for table_idx, table in enumerate(doc.tables):
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                text = clean_text(cell.text)
                if text:
                    cell_texts.append(text)
        if cell_texts:
            blocks.append(
                {
                    "source_file": str(path),
                    "kind": "table",
                    "index": table_idx,
                    "style": None,
                    "text": " | ".join(cell_texts),
                }
            )
    return blocks


def pdf_text_blocks(path: Path) -> list[dict]:
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            text = clean_text(page.extract_text() or "")
            if text:
                blocks.append(
                    {
                        "source_file": str(path),
                        "kind": "pdf_page",
                        "index": page_idx,
                        "style": None,
                        "text": text,
                    }
                )
    return blocks


def text_blocks(paths: list[Path | None]) -> list[dict]:
    blocks = []
    for path in paths:
        if path is None or not path.exists():
            continue
        if path.suffix.lower() == ".docx":
            blocks.extend(docx_text_blocks(path))
        elif path.suffix.lower() == ".pdf":
            blocks.extend(pdf_text_blocks(path))
    return blocks


def extract_gliner_entities(
    blocks: list[dict],
    models: ExtractionModels,
    labels: list[str] = GLINER_LABELS_DE,
    threshold: float = 0.35,
) -> list[dict]:
    if models.gliner_model is None:
        return []

    relevant_blocks = [
        block
        for block in blocks
        if any(term in block["text"].lower() for term in RELEVANT_TERMS)
    ]

    entities = []
    for block in relevant_blocks:
        predictions = models.gliner_model.predict_entities(block["text"][:3500], labels, threshold=threshold)
        for pred in predictions:
            entities.append(
                {
                    "label": pred["label"],
                    "text": clean_text(pred["text"]),
                    "score": round(float(pred["score"]), 3),
                    "source_file": block["source_file"],
                    "block_kind": block["kind"],
                    "block_index": block["index"],
                }
            )
    return entities


def best_text_blocks(
    german_query: str,
    blocks: list[dict],
    models: ExtractionModels,
    top_k: int = 3,
    min_score: float = 0.35,
) -> list[dict]:
    if not blocks:
        return []
    texts = [block["text"] for block in blocks]
    scores = models.cosine_scores(german_query, texts)
    ranked = sorted(enumerate(scores), key=lambda item: float(item[1]), reverse=True)
    results = []
    for idx, score in ranked[:top_k]:
        score = float(score)
        if score >= min_score:
            block = blocks[idx]
            results.append(
                {
                    "score": round(score, 3),
                    "source_file": block["source_file"],
                    "block_kind": block["kind"],
                    "block_index": block["index"],
                    "text": block["text"],
                }
            )
    return results


def top_entities(entities: list[dict], label: str, n: int = 5) -> list[str]:
    rows = [entity for entity in entities if entity["label"] == label]
    rows = sorted(rows, key=lambda item: item["score"], reverse=True)
    return [row["text"] for row in rows[:n]]
